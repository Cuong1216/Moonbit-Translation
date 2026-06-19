import os
import pytest
import docx
import fitz
import tempfile
from services.file_extractor import file_extractor

@pytest.mark.asyncio
async def test_extract_txt(setup_test_fixtures):
    txt_path = setup_test_fixtures["txt"]
    content = await file_extractor.extract_txt(txt_path)
    assert isinstance(content, str)
    assert "Hello World TXT" in content

@pytest.mark.asyncio
async def test_extract_docx(setup_test_fixtures):
    docx_path = setup_test_fixtures["docx"]
    content = await file_extractor.extract_docx(docx_path)
    assert isinstance(content, str)
    assert "Hello World DOCX" in content

@pytest.mark.asyncio
async def test_extract_epub(setup_test_fixtures):
    epub_path = setup_test_fixtures["epub"]
    content = await file_extractor.extract_epub(epub_path)
    assert isinstance(content, str)
    assert "Hello World EPUB" in content

@pytest.mark.asyncio
async def test_extract_pdf(setup_test_fixtures):
    pdf_path = setup_test_fixtures["pdf"]
    content = await file_extractor.extract_pdf(pdf_path)
    assert isinstance(content, str)
    assert "Hello World PDF" in content

@pytest.mark.asyncio
async def test_extract_text_dispatcher(setup_test_fixtures):
    for ext, path in setup_test_fixtures.items():
        res = await file_extractor.extract_text(path, ext)
        assert isinstance(res, dict)
        assert "text" in res
        assert "direction" in res
        assert "toc" in res
        assert isinstance(res["text"], str)
        assert f"Hello World {ext.upper()}" in res["text"]

@pytest.mark.asyncio
async def test_extract_unsupported_format():
    with pytest.raises(ValueError, match="Định dạng tệp không được hỗ trợ"):
        await file_extractor.extract_text("dummy.xyz", ".xyz")

def test_direction_detection_horizontal(setup_test_fixtures):
    # Standard dummy files should be horizontal
    assert file_extractor.detect_text_direction(setup_test_fixtures["txt"]) == "horizontal"
    assert file_extractor.detect_text_direction(setup_test_fixtures["docx"]) == "horizontal"
    assert file_extractor.detect_text_direction(setup_test_fixtures["pdf"]) == "horizontal"
    assert file_extractor.detect_text_direction(setup_test_fixtures["epub"]) == "horizontal"

def test_docx_toc_extraction():
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "test_toc.docx")
        doc = docx.Document()
        doc.add_heading("Chương 1: Mở Đầu", level=1)
        doc.add_paragraph("Nội dung chương 1...")
        doc.add_heading("Chương 1.1: Giới Thiệu Chi Tiết", level=2)
        doc.save(docx_path)
        
        toc = file_extractor.extract_toc(docx_path)
        assert len(toc) == 2
        assert toc[0]["title"] == "Chương 1: Mở Đầu"
        assert toc[0]["level"] == 1
        assert toc[1]["title"] == "Chương 1.1: Giới Thiệu Chi Tiết"
        assert toc[1]["level"] == 2

def test_pdf_toc_extraction():
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "test_toc.pdf")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((100, 100), "Hello PDF TOC")
        # Set outline bookmarks: [level, title, page]
        doc.set_toc([
            [1, "Mục Lục Chương I", 1],
            [2, "Tiểu mục I.A", 1]
        ])
        doc.save(pdf_path)
        doc.close()
        
        toc = file_extractor.extract_toc(pdf_path)
        assert len(toc) == 2
        assert toc[0]["title"] == "Mục Lục Chương I"
        assert toc[0]["level"] == 1
        assert toc[0]["page"] == 1
        assert toc[1]["title"] == "Tiểu mục I.A"
        assert toc[1]["level"] == 2
        assert toc[1]["page"] == 1

@pytest.mark.asyncio
async def test_image_extraction_placeholders(setup_test_fixtures):
    # Verify that extract_text dispatcher generates image directory structure and placeholders
    # (Since setup fixtures have no images, it should return clean text with empty images folder)
    session_id = "test-extract-session-999"
    res = await file_extractor.extract_text(setup_test_fixtures["docx"], ".docx", session_id=session_id)
    assert isinstance(res, dict)
    
    # Check that the temp folder has been created
    temp_dir = os.path.join("static", "temp_images", session_id)
    assert os.path.exists(temp_dir)
    
    # Cleanup temp folder
    if os.path.exists(temp_dir):
        for f in os.listdir(temp_dir):
            os.remove(os.path.join(temp_dir, f))
        os.rmdir(temp_dir)
