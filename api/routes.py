import os
import json
import docx
import tempfile
import asyncio
import logging
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, status, Depends
from fastapi.responses import StreamingResponse, FileResponse
from starlette.background import BackgroundTask
from pydantic import BaseModel
from typing import Dict, Optional, List

# Import các service nghiệp vụ
from services import file_extractor, glossary_builder, novel_translator

# Import DB
from sqlalchemy.orm import Session
from database import get_db, GlossaryCollection, GlossaryTerm

# Cấu hình logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("APIRoutes")

router = APIRouter(prefix="/api", tags=["API"])

# Trạng thái tiến trình dịch đa session
active_sessions = {}

def get_state_file(session_id: str) -> str:
    safe_id = "".join([c for c in session_id if c.isalnum() or c in ('-', '_')])
    return f"state_{safe_id}.json"

def save_session_state(session_id: str):
    session = active_sessions.get(session_id)
    if not session:
        return
    try:
        state_data = {
            "task_id": session_id,
            "chunks": session["chunks"],
            "translated_chunks": session["translated_chunks"],
            "current_chunk_idx": session["current_chunk_idx"],
            "total_chunks": session["total_chunks"],
            "is_paused": session["is_paused"],
            "glossary": session["glossary"],
            "model": session["model"],
            "source_lang": session["source_lang"],
            "target_lang": session["target_lang"],
            "provider": session["provider"]
        }
        with open(get_state_file(session_id), "w", encoding="utf-8") as f:
            json.dump(state_data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Lỗi lưu file trạng thái session {session_id}: {e}")

def delete_session_state(session_id: str):
    file_path = get_state_file(session_id)
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception:
            pass

def load_all_session_states():
    global active_sessions
    try:
        for file_name in os.listdir("."):
            if file_name.startswith("state_") and file_name.endswith(".json"):
                session_id = file_name.replace("state_", "").replace(".json", "")
                try:
                    with open(file_name, "r", encoding="utf-8") as f:
                        data = json.load(f)
                        active_sessions[session_id] = {
                            "chunks": data["chunks"],
                            "translated_chunks": data["translated_chunks"],
                            "current_chunk_idx": data["current_chunk_idx"],
                            "total_chunks": data["total_chunks"],
                            "is_paused": data["is_paused"],
                            "glossary": data["glossary"],
                            "api_key": None,  # Sẽ được nạp lại khi Resume
                            "model": data["model"],
                            "source_lang": data["source_lang"],
                            "target_lang": data["target_lang"],
                            "provider": data["provider"],
                            "queue": asyncio.Queue()
                        }
                    logger.info(f"Đã khôi phục session {session_id} từ file trạng thái.")
                except Exception as e:
                    logger.error(f"Lỗi khôi phục session {session_id}: {e}")
    except Exception as e:
        logger.error(f"Lỗi khi quét file trạng thái session: {e}")

# Khôi phục trạng thái các session khi module được tải
load_all_session_states()

# Pydantic models
class TranslateRequest(BaseModel):
    text: str
    glossary: Optional[Dict[str, str]] = None
    glossary_id: Optional[int] = None
    api_key: Optional[str] = None
    model: str = "gemini-1.5-flash"
    temperature: float = 0.3
    source_lang: str = "Trung Quốc"
    target_lang: str = "Tiếng Việt"
    provider: str = "gemini"

class DownloadRequest(BaseModel):
    text: str
    format: str = "txt"  # "txt" hoặc "docx"
    filename: str = "dich_truyen"

class ExportRequest(BaseModel):
    text: str
    format: str = "txt"  # "txt" hoặc "docx"
    filename: str = "ban_dich"

class ResumeRequest(BaseModel):
    api_key: Optional[str] = None

class CollectionCreate(BaseModel):
    name: str
    description: Optional[str] = None

class TermItem(BaseModel):
    source_term: str
    target_term: str

class TermsUpdate(BaseModel):
    terms: List[TermItem]


@router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """
    Nhận file từ người dùng, gọi FileExtractor để trích xuất text thuần 
    và trả về nội dung cho Frontend hiển thị.
    """
    filename = file.filename
    suffix = os.path.splitext(filename)[1].lower()
    
    if suffix not in [".docx", ".pdf", ".epub", ".txt"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Định dạng tệp không được hỗ trợ. Chỉ hỗ trợ .docx, .pdf, .epub, .txt"
        )
        
    # Lưu tạm tệp tin vào đĩa để FileExtractor xử lý đường dẫn tệp
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        try:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Lỗi ghi file tạm thời: {str(e)}"
            )

    try:
        # Gọi extractor để lấy plain text
        extracted_text = file_extractor.extract_text(tmp_path, suffix)
        return {
            "filename": filename,
            "status": "success",
            "content": extracted_text
        }
    except Exception as e:
        logger.error(f"Lỗi trích xuất file {filename}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Lỗi trích xuất nội dung file: {str(e)}"
        )
    finally:
        # Xóa file tạm thời sau khi xử lý xong
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


@router.post("/build-glossary")
async def build_glossary_endpoint(
    source_file: UploadFile = File(...),
    translated_file: UploadFile = File(...),
    api_key: Optional[str] = Form(None)
):
    """
    Nhận 2 file (gốc và dịch cũ), trích xuất văn bản,
    gọi glossary_builder để đối chiếu và trả về bảng thuật ngữ.
    """
    key = api_key or os.getenv("GEMINI_API_KEY")
    if not key:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Không tìm thấy Gemini API Key. Vui lòng thiết lập trong Cài đặt."
        )

    # Đọc tệp gốc cũ
    src_suffix = os.path.splitext(source_file.filename)[1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=src_suffix) as src_tmp:
        src_tmp.write(await source_file.read())
        src_path = src_tmp.name

    # Đọc tệp dịch cũ
    tgt_suffix = os.path.splitext(translated_file.filename)[1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=tgt_suffix) as tgt_tmp:
        tgt_tmp.write(await translated_file.read())
        tgt_path = tgt_tmp.name

    try:
        # Trích xuất text từ cả hai tệp
        source_text = file_extractor.extract_text(src_path, src_suffix)
        translated_text = file_extractor.extract_text(tgt_path, tgt_suffix)

        if not source_text.strip() or not translated_text.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Một hoặc cả hai tệp tin tải lên không chứa nội dung văn bản."
            )

        # Chạy glossary_builder
        glossary = glossary_builder.build_glossary(
            source_text=source_text,
            translated_text=translated_text,
            api_key=key
        )
        
        return {
            "status": "success",
            "glossary": glossary
        }
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Lỗi xây dựng thuật ngữ: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Lỗi hệ thống khi phân tích thuật ngữ: {str(e)}"
        )
    finally:
        # Dọn dẹp file tạm
        for path in [src_path, tgt_path]:
            if os.path.exists(path):
                os.remove(path)


async def run_translation_task(session_id: str):
    session = active_sessions.get(session_id)
    if not session:
        return
    
    total = session["total_chunks"]
    queue = session["queue"]
    
    while session["current_chunk_idx"] < total:
        if session["is_paused"]:
            await queue.put({
                "status": "Đã tạm dừng dịch thuật.",
                "progress": int((session["current_chunk_idx"] / total) * 100),
                "event": "paused",
                "current_chunk_idx": session["current_chunk_idx"]
            })
            return
            
        idx = session["current_chunk_idx"]
        chunk = session["chunks"][idx]
        progress = int((idx / total) * 100)
        
        await queue.put({
            "status": f"Đang dịch đoạn {idx + 1}/{total}...",
            "progress": progress,
            "event": "translating"
        })
        
        try:
            translation = await asyncio.to_thread(
                novel_translator.translate_chunk,
                chunk,
                session["glossary"],
                session["api_key"],
                session["model"],
                session["source_lang"],
                session["target_lang"],
                session["provider"]
            )
            
            if session["is_paused"]:
                session["translated_chunks"].append(translation)
                session["current_chunk_idx"] += 1
                save_session_state(session_id)
                
                await queue.put({
                    "status": f"Đoạn {idx + 1} đã dịch xong trước khi tạm dừng.",
                    "progress": int(((idx + 1) / total) * 100),
                    "event": "chunk_completed",
                    "source_chunk": chunk,
                    "translated_chunk": translation
                })
                
                await queue.put({
                    "status": "Đã tạm dừng dịch thuật.",
                    "progress": int(((idx + 1) / total) * 100),
                    "event": "paused",
                    "current_chunk_idx": idx + 1
                })
                return
            
            session["translated_chunks"].append(translation)
            session["current_chunk_idx"] += 1
            save_session_state(session_id)
            
            await queue.put({
                "status": f"Hoàn thành đoạn {idx + 1}/{total}",
                "progress": int(((idx + 1) / total) * 100),
                "event": "chunk_completed",
                "source_chunk": chunk,
                "translated_chunk": translation
            })
            
        except Exception as e:
            logger.error(f"Lỗi khi dịch đoạn {idx + 1} trong session {session_id}: {e}")
            await queue.put({
                "status": f"Thất bại ở đoạn {idx + 1}: {str(e)}",
                "progress": progress,
                "event": "error",
                "error": str(e)
            })
            return

    full_text = "\n\n".join(session["translated_chunks"])
    delete_session_state(session_id)
    
    await queue.put({
        "status": "Dịch hoàn tất!",
        "progress": 100,
        "event": "completed",
        "full_translation": full_text
    })

async def event_generator(session_id: str):
    session = active_sessions.get(session_id)
    if not session:
        yield f"data: {json.dumps({'status': 'Session không tồn tại.', 'event': 'error', 'error': 'Invalid session ID'}, ensure_ascii=False)}\n\n"
        return
        
    queue = session["queue"]
    
    # Gửi sự kiện phục hồi (restore_state) cho client nếu có tiến trình trước đó
    if session["current_chunk_idx"] > 0:
        restore_data = {
            'status': f'Khôi phục tiến trình dịch: {session["current_chunk_idx"]}/{session["total_chunks"]}',
            'progress': int((session["current_chunk_idx"] / session["total_chunks"]) * 100),
            'event': 'restore_state',
            'current_chunk_idx': session["current_chunk_idx"],
            'total_chunks': session["total_chunks"],
            'chunks': session["chunks"][:session["current_chunk_idx"]],
            'translated_chunks': session["translated_chunks"]
        }
        yield f"data: {json.dumps(restore_data, ensure_ascii=False)}\n\n"
    
    while True:
        try:
            event = await queue.get()
            yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
            if event["event"] in ["completed", "paused", "error"]:
                break
        except Exception as e:
            logger.error(f"Lỗi trong event generator cho session {session_id}: {e}")
            break

@router.post("/translate")
async def translate_endpoint(request: TranslateRequest, db: Session = Depends(get_db)):
    """
    Tạo một phiên dịch thuật mới, trả về session_id lập tức và chạy ngầm tiến trình dịch.
    """
    provider = request.provider.strip().lower()
    
    key = request.api_key
    if not key:
        if provider == "gemini":
            key = os.getenv("GEMINI_API_KEY")
        elif provider == "claude":
            key = os.getenv("ANTHROPIC_API_KEY")
        elif provider == "openrouter":
            key = os.getenv("OPENROUTER_API_KEY")

    if not key:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Không tìm thấy API Key cho {request.provider}. Vui lòng thiết lập trong Cài đặt."
        )

    chunks = novel_translator.smart_chunking(request.text)
    if not chunks:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Văn bản nguồn rỗng hoặc không thể chia nhỏ."
        )

    glossary_dict = {}
    if request.glossary_id is not None:
        terms = db.query(GlossaryTerm).filter(GlossaryTerm.collection_id == request.glossary_id).all()
        glossary_dict = {t.source_term: t.target_term for t in terms}
    elif request.glossary:
        glossary_dict = request.glossary

    import uuid
    session_id = str(uuid.uuid4())
    
    active_sessions[session_id] = {
        "chunks": chunks,
        "translated_chunks": [],
        "current_chunk_idx": 0,
        "total_chunks": len(chunks),
        "is_paused": False,
        "glossary": glossary_dict,
        "api_key": key,
        "model": request.model,
        "source_lang": request.source_lang,
        "target_lang": request.target_lang,
        "provider": provider,
        "queue": asyncio.Queue()
    }
    save_session_state(session_id)

    asyncio.create_task(run_translation_task(session_id))

    return {"status": "started", "session_id": session_id}

@router.get("/stream")
async def stream_endpoint(session_id: str):
    """
    SSE stream endpoint để Frontend kết nối nhận tiến trình dịch thời gian thực.
    """
    if session_id not in active_sessions:
        raise HTTPException(status_code=404, detail="Không tìm thấy phiên dịch thuật này.")
    return StreamingResponse(event_generator(session_id), media_type="text/event-stream")

@router.post("/pause")
async def pause_endpoint(session_id: str):
    session = active_sessions.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Không tìm thấy phiên dịch thuật này.")
    
    session["is_paused"] = True
    save_session_state(session_id)
    return {
        "status": "success",
        "message": "Đã gửi yêu cầu tạm dừng.",
        "current_chunk_idx": session["current_chunk_idx"],
        "total_chunks": session["total_chunks"]
    }

@router.post("/resume")
async def resume_endpoint(session_id: str, request: Optional[ResumeRequest] = None):
    session = active_sessions.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Không tìm thấy phiên dịch thuật để tiếp tục.")
    
    if not session["is_paused"]:
        return {"status": "ignored", "message": "Tiến trình dịch đang chạy, không cần resume."}
        
    session["is_paused"] = False
    if request and request.api_key:
        session["api_key"] = request.api_key
        
    session["queue"] = asyncio.Queue()
    save_session_state(session_id)
    
    asyncio.create_task(run_translation_task(session_id))
    return {"status": "resumed", "session_id": session_id}


@router.post("/download")
async def download_file_endpoint(request: DownloadRequest):
    """
    Cho phép tải file text hoặc docx chứa kết quả sau khi dịch xong.
    Sử dụng BackgroundTask để dọn dẹp tệp tạm thời sau khi tải xong.
    """
    if not request.text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Nội dung tải xuống rỗng."
        )

    # Đảm bảo tên file an toàn
    safe_filename = "".join([c for c in request.filename if c.isalpha() or c.isdigit() or c in (' ', '_', '-')]).rstrip()
    if not safe_filename:
        safe_filename = "ban_dich"

    if request.format == "docx":
        # Tạo file Word docx
        doc = docx.Document()
        
        # Chia các đoạn văn theo dấu xuống dòng và ghi vào tệp docx
        for para in request.text.split("\n"):
            doc.add_paragraph(para)
            
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp_path = tmp.name
        tmp.close()
        
        try:
            doc.save(tmp_path)
            
            # Đăng ký BackgroundTask xóa file tạm sau khi response được gửi xong
            cleanup_task = BackgroundTask(lambda: os.path.exists(tmp_path) and os.remove(tmp_path))
            
            return FileResponse(
                tmp_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"{safe_filename}.docx",
                background=cleanup_task
            )
        except Exception as e:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Lỗi tạo tệp tin Word: {str(e)}"
            )
    else:
        # Mặc định tải về file .txt
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8")
        tmp_path = tmp.name
        
        try:
            tmp.write(request.text)
            tmp.close()
            
            cleanup_task = BackgroundTask(lambda: os.path.exists(tmp_path) and os.remove(tmp_path))
            
            return FileResponse(
                tmp_path,
                media_type="text/plain",
                filename=f"{safe_filename}.txt",
                background=cleanup_task
            )
        except Exception as e:
            tmp.close()
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Lỗi tạo tệp tin Text: {str(e)}"
            )

@router.post("/export")
async def export_file_endpoint(request: ExportRequest):
    """
    Tương tự như download_file_endpoint, nhận nội dung chỉnh sửa từ Frontend
    và xuất tệp .txt hoặc .docx để người dùng tải về.
    """
    if not request.text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Nội dung xuất file rỗng."
        )

    # Đảm bảo tên file an toàn
    safe_filename = "".join([c for c in request.filename if c.isalpha() or c.isdigit() or c in (' ', '_', '-')]).rstrip()
    if not safe_filename:
        safe_filename = "ban_dich"

    if request.format == "docx":
        doc = docx.Document()
        for para in request.text.split("\n"):
            doc.add_paragraph(para)
            
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp_path = tmp.name
        tmp.close()
        
        try:
            doc.save(tmp_path)
            cleanup_task = BackgroundTask(lambda: os.path.exists(tmp_path) and os.remove(tmp_path))
            return FileResponse(
                tmp_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"{safe_filename}.docx",
                background=cleanup_task
            )
        except Exception as e:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Lỗi tạo tệp tin Word: {str(e)}"
            )
    else:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8")
        tmp_path = tmp.name
        try:
            tmp.write(request.text)
            tmp.close()
            cleanup_task = BackgroundTask(lambda: os.path.exists(tmp_path) and os.remove(tmp_path))
            return FileResponse(
                tmp_path,
                media_type="text/plain",
                filename=f"{safe_filename}.txt",
                background=cleanup_task
            )
        except Exception as e:
            tmp.close()
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Lỗi tạo tệp tin Text: {str(e)}"
            )

@router.get("/glossaries")
def get_glossaries(db: Session = Depends(get_db)):
    collections = db.query(GlossaryCollection).all()
    result = []
    for col in collections:
        count = db.query(GlossaryTerm).filter(GlossaryTerm.collection_id == col.id).count()
        result.append({
            "id": col.id,
            "name": col.name,
            "description": col.description,
            "created_at": col.created_at.isoformat() if col.created_at else None,
            "term_count": count
        })
    return result

@router.post("/glossaries")
def create_glossary(data: CollectionCreate, db: Session = Depends(get_db)):
    name_clean = data.name.strip()
    if not name_clean:
        raise HTTPException(status_code=400, detail="Tên bộ từ vựng không được để trống.")
    
    existing = db.query(GlossaryCollection).filter(GlossaryCollection.name == name_clean).first()
    if existing:
        raise HTTPException(status_code=400, detail="Bộ từ vựng với tên này đã tồn tại.")
        
    col = GlossaryCollection(name=name_clean, description=data.description)
    db.add(col)
    db.commit()
    db.refresh(col)
    return {
        "status": "success", 
        "collection": {
            "id": col.id, 
            "name": col.name, 
            "description": col.description,
            "term_count": 0
        }
    }

@router.delete("/glossaries/{id}")
def delete_glossary(id: int, db: Session = Depends(get_db)):
    col = db.query(GlossaryCollection).filter(GlossaryCollection.id == id).first()
    if not col:
        raise HTTPException(status_code=404, detail="Không tìm thấy bộ từ vựng này.")
    db.delete(col)
    db.commit()
    return {"status": "success", "message": "Đã xóa bộ từ vựng thành công."}

@router.get("/glossaries/{id}/terms")
def get_glossary_terms(id: int, db: Session = Depends(get_db)):
    col = db.query(GlossaryCollection).filter(GlossaryCollection.id == id).first()
    if not col:
        raise HTTPException(status_code=404, detail="Không tìm thấy bộ từ vựng này.")
    
    terms = db.query(GlossaryTerm).filter(GlossaryTerm.collection_id == id).all()
    result = {t.source_term: t.target_term for t in terms}
    return {"status": "success", "terms": result}

@router.post("/glossaries/{id}/terms")
def update_glossary_terms(id: int, request: TermsUpdate, db: Session = Depends(get_db)):
    col = db.query(GlossaryCollection).filter(GlossaryCollection.id == id).first()
    if not col:
        raise HTTPException(status_code=404, detail="Không tìm thấy bộ từ vựng này.")
    
    # Xóa các terms cũ của bộ từ vựng này
    db.query(GlossaryTerm).filter(GlossaryTerm.collection_id == id).delete()
    
    # Thêm các terms mới
    new_terms = []
    for t in request.terms:
        src = t.source_term.strip()
        tgt = t.target_term.strip()
        if src and tgt:
            new_terms.append(GlossaryTerm(collection_id=id, source_term=src, target_term=tgt))
            
    if new_terms:
        db.bulk_save_objects(new_terms)
        
    db.commit()
    return {"status": "success", "message": f"Đã lưu thành công {len(new_terms)} thuật ngữ."}
