import os
import asyncio
import docx
import pdfplumber
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import warnings
import uuid

# Tắt các cảnh báo không cần thiết từ EbookLib (như cảnh báo về chuẩn XML / Dublin Core)
warnings.filterwarnings("ignore", category=UserWarning, module="ebooklib")
warnings.filterwarnings("ignore", category=FutureWarning)

class FileExtractor:
    """
    Lớp tiện ích trích xuất nội dung văn bản thuần (plain text) từ các định dạng tệp khác nhau:
    - .docx (Microsoft Word)
    - .pdf (Portable Document Format)
    - .epub (Electronic Publication - Sách điện tử)
    - .txt (Tệp văn bản thuần)
    
    Hỗ trợ xử lý nâng cao: trích xuất ảnh, TOC, tự động nhận diện hướng chữ.
    """

    def detect_text_direction(self, file_path: str) -> str:
        """
        Tự động nhận diện hướng chữ (Layout Detection):
        Trả về 'vertical' (nếu là chữ dọc tiếng Nhật) hoặc 'horizontal' (nếu là chữ ngang).
        Dựa vào CSS writing-mode (EPUB), thuộc tính w:textDirection trong XML (DOCX), hoặc wmode/dir (PDF).
        """
        if not os.path.exists(file_path):
            return 'horizontal'
            
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.epub':
            try:
                book = epub.read_epub(file_path)
                for item in book.get_items():
                    if item.get_type() in [ebooklib.ITEM_STYLE, ebooklib.ITEM_DOCUMENT]:
                        content = item.get_content().decode('utf-8', errors='ignore')
                        if 'writing-mode' in content:
                            if any(val in content for val in ['vertical-rl', 'vertical-lr', 'epub-vertical']):
                                return 'vertical'
            except Exception:
                pass
        elif ext == '.docx':
            try:
                import zipfile
                with zipfile.ZipFile(file_path) as z:
                    for name in z.namelist():
                        if name.endswith('.xml'):
                            xml_content = z.read(name).decode('utf-8', errors='ignore')
                            if 'textDirection' in xml_content or 'w:textDirection' in xml_content:
                                if any(val in xml_content for val in ['tbRl', 'tbRlV', 'vertical', 'btLr']):
                                    return 'vertical'
            except Exception:
                pass
        elif ext == '.pdf':
            try:
                import fitz
                with fitz.open(file_path) as doc:
                    vertical_count = 0
                    horizontal_count = 0
                    for page in doc:
                        blocks = page.get_text("dict").get("blocks", [])
                        for b in blocks:
                            if "lines" in b:
                                for l in b["lines"]:
                                    if l.get("wmode") == 1 or (l.get("dir") and abs(l.get("dir")[1]) > abs(l.get("dir")[0])):
                                        vertical_count += 1
                                    else:
                                        horizontal_count += 1
                    if vertical_count > horizontal_count and vertical_count > 0:
                        return 'vertical'
            except Exception:
                pass
        return 'horizontal'

    def _parse_epub_toc(self, toc, level=1) -> list:
        result = []
        for item in toc:
            if isinstance(item, tuple):
                parent, children = item
                if hasattr(parent, 'title') and hasattr(parent, 'href'):
                    result.append({
                        "title": parent.title,
                        "level": level,
                        "href": parent.href
                    })
                result.extend(self._parse_epub_toc(children, level + 1))
            elif hasattr(item, 'title') and hasattr(item, 'href'):
                result.append({
                    "title": item.title,
                    "level": level,
                    "href": item.href
                })
        return result

    def _parse_docx_toc(self, doc) -> list:
        result = []
        for p in doc.paragraphs:
            style_name = p.style.name if p.style else ""
            if style_name and style_name.startswith("Heading"):
                try:
                    level_str = ''.join(filter(str.isdigit, style_name))
                    level = int(level_str) if level_str else 1
                except Exception:
                    level = 1
                
                if p.text.strip():
                    result.append({
                        "title": p.text.strip(),
                        "level": level
                    })
        return result

    def _parse_pdf_toc(self, doc) -> list:
        toc = doc.get_toc()
        result = []
        for item in toc:
            if len(item) >= 3:
                result.append({
                    "title": item[1],
                    "level": item[0],
                    "page": item[2]
                })
        return result

    def extract_toc(self, file_path: str) -> list:
        """
        Trích xuất cấu trúc mục lục (nếu có) thành một mảng JSON riêng để xử lý độc lập.
        """
        if not os.path.exists(file_path):
            return []
            
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.epub':
            try:
                book = epub.read_epub(file_path)
                return self._parse_epub_toc(book.toc)
            except Exception:
                return []
        elif ext == '.pdf':
            try:
                import fitz
                with fitz.open(file_path) as doc:
                    return self._parse_pdf_toc(doc)
            except Exception:
                return []
        elif ext == '.docx':
            try:
                doc = docx.Document(file_path)
                return self._parse_docx_toc(doc)
            except Exception:
                return []
        return []

    async def extract_docx(self, file_path: str, session_id: str = None) -> str:
        """
        Trích xuất văn bản từ tệp DOCX bằng thư viện python-docx.
        Đọc từng paragraph và hỗ trợ bóc tách ảnh lưu vào static/temp_images/[session_id]/.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Không tìm thấy tệp DOCX tại: {file_path}")
        
        session_id = session_id or str(uuid.uuid4())
        temp_dir = os.path.join("static", "temp_images", session_id)
        os.makedirs(temp_dir, exist_ok=True)
        
        def _extract():
            doc = docx.Document(file_path)
            paragraphs_text = []
            
            for p in doc.paragraphs:
                p_text = ""
                has_images = False
                for run in p.runs:
                    run_text = run.text
                    try:
                        rIds = run._r.xpath('.//a:blip/@r:embed')
                    except Exception:
                        rIds = []
                    
                    image_placeholders = []
                    for rId in rIds:
                        try:
                            if rId in doc.part.related_parts:
                                image_part = doc.part.related_parts[rId]
                                img_basename = os.path.basename(image_part.partname)
                                image_name = f"docx_{img_basename}"
                                img_path = os.path.join(temp_dir, image_name)
                                with open(img_path, 'wb') as f:
                                    f.write(image_part.blob)
                                image_placeholders.append(f"---[IMAGE:{image_name}]---")
                                has_images = True
                        except Exception:
                            pass
                    
                    if image_placeholders:
                        p_text += run_text + " " + " ".join(image_placeholders) + " "
                    else:
                        p_text += run_text
                
                if not has_images and not p_text.strip() and p.text.strip():
                    p_text = p.text
                
                paragraphs_text.append(p_text.strip())
                
            return "\n".join(paragraphs_text)
            
        return await asyncio.to_thread(_extract)

    async def extract_pdf(self, file_path: str, session_id: str = None) -> str:
        """
        Trích xuất văn bản từ tệp PDF bằng PyMuPDF (fitz) có sắp xếp theo hướng chữ,
        và bóc tách hình ảnh lưu vào static/temp_images/[session_id]/.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Không tìm thấy tệp PDF tại: {file_path}")
            
        session_id = session_id or str(uuid.uuid4())
        temp_dir = os.path.join("static", "temp_images", session_id)
        os.makedirs(temp_dir, exist_ok=True)
        
        def _extract():
            try:
                import fitz
                doc = fitz.open(file_path)
                direction = self.detect_text_direction(file_path)
                
                text_pages = []
                for page_idx, page in enumerate(doc):
                    raw_blocks = page.get_text("blocks")
                    elements = []
                    for b in raw_blocks:
                        elements.append({
                            "type": "text",
                            "bbox": b[0:4],
                            "content": b[4]
                        })
                    
                    image_list = page.get_images(full=True)
                    for img_idx, img in enumerate(image_list):
                        xref = img[0]
                        rects = page.get_image_rects(xref)
                        if rects:
                            bbox = (rects[0].x0, rects[0].y0, rects[0].x1, rects[0].y1)
                        else:
                            bbox = (0, 0, 0, 0)
                        
                        try:
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            image_ext = base_image["ext"]
                            image_name = f"pdf_p{page_idx}_img{img_idx}_{xref}.{image_ext}"
                            img_path = os.path.join(temp_dir, image_name)
                            with open(img_path, "wb") as f:
                                f.write(image_bytes)
                            
                            elements.append({
                                "type": "image",
                                "bbox": bbox,
                                "content": f"\n---[IMAGE:{image_name}]---\n"
                            })
                        except Exception:
                            pass
                    
                    if direction == "vertical":
                        elements.sort(key=lambda e: (-e["bbox"][2], e["bbox"][1]))
                    else:
                        elements.sort(key=lambda e: (e["bbox"][1], e["bbox"][0]))
                    
                    page_text = ""
                    for el in elements:
                        page_text += el["content"]
                    
                    text_pages.append(page_text.strip())
                    
                return "\n\n".join(text_pages)
                
            except Exception as fitz_err:
                # Fallback to pdfplumber if fitz has errors
                text_pages = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_pages.append(page_text)
                return "\n\n".join(text_pages)
                
        return await asyncio.to_thread(_extract)

    async def extract_epub(self, file_path: str, session_id: str = None) -> str:
        """
        Trích xuất văn bản từ tệp EPUB bằng ebooklib và BeautifulSoup4.
        Hỗ trợ trích xuất hình ảnh lưu vào thư mục tạm static/temp_images/[session_id]/.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Không tìm thấy tệp EPUB tại: {file_path}")
            
        session_id = session_id or str(uuid.uuid4())
        temp_dir = os.path.join("static", "temp_images", session_id)
        os.makedirs(temp_dir, exist_ok=True)
        
        def _extract():
            book = epub.read_epub(file_path)
            text_parts = []
            
            def find_image_item(src, book):
                src_clean = src.split('?')[0].split('/')[-1]
                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_IMAGE:
                        if item.file_name.endswith(src_clean) or item.get_name().endswith(src_clean):
                            return item
                return None
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    html_content = item.get_content()
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    for img in soup.find_all(['img', 'image']):
                        src = img.get('src') or img.get('href') or img.get('xlink:href')
                        if src:
                            img_item = find_image_item(src, book)
                            if img_item:
                                image_name = img_item.file_name.replace('/', '_').replace('\\', '_')
                                img_path = os.path.join(temp_dir, image_name)
                                try:
                                    with open(img_path, 'wb') as f:
                                        f.write(img_item.get_content())
                                    img.replace_with(f"\n---[IMAGE:{image_name}]---\n")
                                except Exception:
                                    pass
                    
                    text = soup.get_text()
                    cleaned_text = text.strip()
                    if cleaned_text:
                        text_parts.append(cleaned_text)
            return "\n\n".join(text_parts)
            
        return await asyncio.to_thread(_extract)

    async def extract_txt(self, file_path: str) -> str:
        """
        Đọc tệp văn bản thuần .txt với các bảng mã phổ biến (utf-8, utf-8-sig, latin-1).
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Không tìm thấy tệp TXT tại: {file_path}")
            
        def _extract():
            encodings = ['utf-8', 'utf-8-sig', 'utf-16', 'ansi', 'latin-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        return await asyncio.to_thread(_extract)

    async def extract_text(self, file_path: str, extension: str, session_id: str = None) -> dict:
        """
        Hàm chính tự động phát hiện định dạng tệp và trả về dictionary:
        {"text": text, "direction": direction, "toc": toc}.
        """
        ext = extension.strip().lower()
        if not ext.startswith('.'):
            ext = '.' + ext
            
        # Detect direction & TOC
        direction = self.detect_text_direction(file_path)
        toc = self.extract_toc(file_path)
        
        if ext == '.docx':
            text = await self.extract_docx(file_path, session_id=session_id)
        elif ext == '.pdf':
            text = await self.extract_pdf(file_path, session_id=session_id)
        elif ext == '.epub':
            text = await self.extract_epub(file_path, session_id=session_id)
        elif ext == '.txt':
            text = await self.extract_txt(file_path)
        else:
            raise ValueError(f"Định dạng tệp không được hỗ trợ: {extension}. Chỉ hỗ trợ .docx, .pdf, .epub, .txt")
            
        return {
            "text": text,
            "direction": direction,
            "toc": toc
        }

# Khởi tạo một instance mặc định để sử dụng nhanh
file_extractor = FileExtractor()
